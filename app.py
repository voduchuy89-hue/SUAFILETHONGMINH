import streamlit as st
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes
import io
import os
from openai import OpenAI
import docx
from docx.shared import Mm, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font
from openpyxl import load_workbook
import re
import json
import zipfile
import unicodedata

# ========================================================================================
# CẤU HÌNH TRANG
# ========================================================================================
MAX_FILES = 20  # Tối đa 20 file cùng lúc

st.set_page_config(
    page_title="Trợ lý OCR Thông minh",
    page_icon="📄",
    layout="wide"
)

# ========================================================================================
# HÀM HỖ TRỢ (LOGIC XỬ LÝ)
# ========================================================================================

@st.cache_data  # Sử dụng cache để không xử lý lại file đã xử lý
def process_file(file_bytes, file_extension, show_progress=True):
    """
    Hàm trung tâm xử lý file đầu vào (ảnh hoặc PDF) và trả về văn bản được trích xuất.
    Mặc định sử dụng chế độ song ngữ Việt + Anh.
    show_progress=False dùng khi xử lý hàng loạt nhiều file.
    """
    lang_code = "vie+eng"
    extracted_text = ""
    try:
        if file_extension == 'pdf':
            images = convert_from_bytes(file_bytes)
            all_text = []
            progress_bar = st.progress(0, text="Đang xử lý file PDF...") if show_progress else None
            for i, img in enumerate(images):
                all_text.append(pytesseract.image_to_string(img, lang=lang_code))
                if progress_bar is not None:
                    progress_bar.progress((i + 1) / len(images))
            extracted_text = "\n\n--- Hết trang ---\n\n".join(all_text)
        elif file_extension in ['png', 'jpg', 'jpeg']:
            image = Image.open(io.BytesIO(file_bytes))
            extracted_text = pytesseract.image_to_string(image, lang=lang_code)
        return extracted_text, None
    except Exception as e:
        return None, f"Đã xảy ra lỗi trong quá trình xử lý: {e}"


def call_openai_proofread(text: str) -> str:
    """
    Gọi OpenAI để kiểm tra lỗi, chỉnh sửa văn bản sau OCR.
    Ưu tiên lấy khóa từ Streamlit secrets (OPENAI_API_KEY).
    """
    api_key = os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY", None)
    if not api_key:
        raise RuntimeError(
            "Chưa cấu hình OpenAI. Vào Manage app → Settings → Secrets và thêm OPENAI_API_KEY."
        )

    # Đảm bảo thư viện OpenAI đọc được API key từ biến môi trường
    os.environ["OPENAI_API_KEY"] = api_key

    client = OpenAI()
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": (
                    "Bạn là trợ lý giúp chỉnh sửa văn bản tiếng Việt và tiếng Anh "
                    "được trích xuất từ OCR. Hãy sửa lỗi chính tả, dấu câu và cách "
                    "xuống dòng hợp lý. Chỉ trả về văn bản đã chỉnh sửa, không giải thích gì thêm."
                ),
            },
            {"role": "user", "content": text},
        ],
        temperature=0.2,
    )
    return response.choices[0].message.content


def build_docx(text: str) -> bytes:
    """
    Tạo file Word (.docx) khổ A4 từ văn bản.
    """
    buffer = io.BytesIO()
    document = docx.Document()

    # Thiết lập khổ giấy A4
    section = document.sections[0]
    section.page_width = Mm(210)   # A4 ngang 210mm
    section.page_height = Mm(297)  # A4 dọc 297mm

    # Thiết lập font mặc định Times New Roman, cỡ 12
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)

    for line in text.splitlines():
        document.add_paragraph(line)

    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_excel(text: str) -> bytes:
    """
    Tạo file Excel (.xlsx) với nội dung ở cột A, thiết lập in trên khổ A4.
    """
    buffer = io.BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = "OCR"

    # Thiết lập khổ giấy A4 khi in
    ws.page_setup.paperSize = ws.PAPERSIZE_A4

    # Font chung cho toàn bộ nội dung
    base_font = Font(name="Times New Roman", size=12)

    lines = text.splitlines()
    for idx, line in enumerate(lines, start=1):
        cell = ws[f"A{idx}"]
        cell.value = line
        cell.alignment = Alignment(wrap_text=True)
        cell.font = base_font

    ws.column_dimensions["A"].width = 100

    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ========================================================================================
# HÀM HỖ TRỢ CHO TEMPLATE FILLING
# ========================================================================================

def extract_text_from_file(file_bytes, file_extension):
    """
    Trích xuất text từ các loại file khác nhau:
    - PDF, ảnh (PNG, JPG): dùng OCR
    - Word (.docx): đọc từ document
    - Excel (.xlsx): đọc từ cells
    - Text (.txt): đọc trực tiếp
    """
    try:
        if file_extension == 'txt':
            return file_bytes.decode('utf-8')
        elif file_extension == 'docx':
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n".join([p.text for p in doc.paragraphs])
            return text
        elif file_extension == 'xlsx':
            wb = load_workbook(io.BytesIO(file_bytes))
            ws = wb.active
            text = []
            for row in ws.iter_rows(values_only=True):
                text.append(" | ".join([str(cell) if cell else "" for cell in row]))
            return "\n".join(text)
        else:
            # Nếu là ảnh hoặc PDF, dùng process_file (OCR)
            extracted_text, err = process_file(file_bytes, file_extension, show_progress=False)
            if err:
                return f"Lỗi OCR: {err}"
            return extracted_text
    except Exception as e:
        return f"Lỗi trích xuất: {str(e)}"


def detect_placeholders_in_template(template_bytes):
    """
    Phát hiện tất cả placeholders dạng {field_name} trong template Word.
    Trả về danh sách các field name cần điền.
    """
    try:
        doc = docx.Document(io.BytesIO(template_bytes))
        placeholders = set()
        pattern = r'\{([^}]+)\}'
        
        # Scan trong paragraphs
        for paragraph in doc.paragraphs:
            matches = re.findall(pattern, paragraph.text)
            placeholders.update(matches)
        
        # Scan trong tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(pattern, paragraph.text)
                        placeholders.update(matches)
        
        return sorted(list(placeholders))
    except Exception as e:
        return []


def fill_template_word(template_bytes, data_dict):
    """
    Điền dữ liệu vào template Word.
    data_dict: {"field_name": "value", ...}
    Trả về bytes của document đã điền.
    """
    try:
        doc = docx.Document(io.BytesIO(template_bytes))
        
        # Hàm helper để thay thế placeholder trong paragraph
        def replace_in_paragraph(paragraph, data_dict):
            # Get text từ tất cả runs
            full_text = "".join([run.text for run in paragraph.runs])
            original_text = full_text
            
            # Thay thế tất cả placeholder
            modified_text = full_text
            for placeholder, value in data_dict.items():
                pattern = r'\{' + re.escape(placeholder) + r'\}'
                modified_text = re.sub(pattern, str(value) if value else "", modified_text)
            
            # Nếu có thay đổi, cập nhật paragraph
            if modified_text != original_text:
                # Xóa tất cả run cũ
                for run in list(paragraph.runs):
                    r = run._element
                    r.getparent().remove(r)
                
                # Thêm text mới
                if modified_text:
                    paragraph.add_run(modified_text)
        
        # Điền trong paragraphs
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, data_dict)
        
        # Điền trong tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, data_dict)
        
        # Lưu vào BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        raise Exception(f"Lỗi điền template: {str(e)}")


def fill_template_with_labels(template_bytes, data_dict):
    """
    Fallback filler: nếu template không có placeholders, tìm các label (key)
    trong văn bản mẫu và điền giá trị sau label.
    - Nếu paragraph chứa chính xác label, thay bằng 'label: value'
    - Nếu paragraph chứa label kèm text khác, sẽ thay phần label bằng 'label: value'
    Trả về bytes của document đã điền.
    """
    try:
        doc = docx.Document(io.BytesIO(template_bytes))

        def replace_label_in_paragraph(paragraph, data_dict):
            para_text = paragraph.text
            modified = para_text
            for key, val in data_dict.items():
                if not key:
                    continue
                # match case-insensitive
                idx = para_text.lower().find(key.lower())
                if idx != -1:
                    # If paragraph is exactly the label or contains only small punctuation after
                    if para_text.strip().lower() == key.lower() or para_text.strip().lower().startswith(key.lower()):
                        # Replace the first occurrence with 'key: value'
                        # Build replacement preserving original label casing
                        start = para_text[:idx]
                        end = para_text[idx+len(key):]
                        replacement = f"{para_text[idx:idx+len(key)]}: {val}"
                        modified = (start + replacement + end).strip()
                    else:
                        # For other cases, replace label occurrence with 'label: value'
                        modified = re.sub(re.escape(key), f"{key}: {val}", modified, flags=re.IGNORECASE)
            if modified != para_text:
                # remove runs
                for run in list(paragraph.runs):
                    r = run._element
                    r.getparent().remove(r)
                paragraph.add_run(modified)

        for paragraph in doc.paragraphs:
            replace_label_in_paragraph(paragraph, data_dict)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_label_in_paragraph(paragraph, data_dict)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        raise Exception(f"Lỗi điền template (label fallback): {str(e)}")


def normalize_text(s: str) -> str:
    """Lowercase, strip, remove accents and non-alphanum for matching."""
    if not s:
        return ""
    s = s.lower().strip()
    s = unicodedata.normalize('NFKD', s)
    s = ''.join(ch for ch in s if not unicodedata.combining(ch))
    # keep alnum and spaces
    s = re.sub(r'[^0-9a-z\s]', ' ', s)
    s = re.sub(r'\s+', ' ', s).strip()
    return s


def detect_labels_in_template(template_bytes):
    """
    Scan a .docx template and return likely label strings.
    Heuristics: paragraphs with short length (<=80 chars) or ending with ':'
    """
    try:
        doc = docx.Document(io.BytesIO(template_bytes))
        labels = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if len(text) <= 80 or text.endswith(':') or '\t' in text:
                # split by colon if present and take left part as label
                if ':' in text:
                    left = text.split(':', 1)[0].strip()
                    if left:
                        labels.append(left)
                        continue
                # if short line, treat as label
                # ignore lines that look like long sentences
                if len(text.split()) <= 6:
                    labels.append(text)
        # deduplicate preserving order
        seen = set()
        out = []
        for l in labels:
            if l not in seen:
                seen.add(l)
                out.append(l)
        return out
    except Exception:
        return []


def map_data_to_labels(data_dict, labels):
    """
    Map AI-extracted keys to template labels using normalized matching.
    Returns mapped dict: {label: value}
    """
    mapped = {label: "" for label in labels}
    if not data_dict or not labels:
        return mapped

    # Precompute normalized keys
    norm_keys = {k: normalize_text(k) for k in data_dict.keys()}
    norm_labels = {lab: normalize_text(lab) for lab in labels}

    for lab, nlab in norm_labels.items():
        best_key = None
        best_score = 0
        for k, nk in norm_keys.items():
            score = 0
            if nk and nlab:
                if nk == nlab:
                    score = 100
                elif nk in nlab or nlab in nk:
                    score = 70
                else:
                    # word overlap
                    set_k = set(nk.split())
                    set_l = set(nlab.split())
                    inter = set_k & set_l
                    if inter:
                        score = 30 + 10 * len(inter)
            if score > best_score:
                best_score = score
                best_key = k
        if best_key and best_score >= 30:
            mapped[lab] = data_dict.get(best_key, "")
    return mapped








def extract_structured_data_with_ai(text, placeholders):
    """
    Dùng OpenAI để trích xuất & lọc dữ liệu có cấu trúc từ text.
    AI sẽ quyết định lấy thông tin nào từ text để phù hợp với placeholders.
    Trả về dict: {"field_name": "value", ...}
    """
    api_key = os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY", None)
    if not api_key:
        raise RuntimeError("Chưa cấu hình OpenAI API Key.")
    
    os.environ["OPENAI_API_KEY"] = api_key
    client = OpenAI()
    
    # Nếu không có placeholders, yêu cầu AI tự quyết định các trường cần trích xuất
    if not placeholders:
        prompt = f"""Bạn là chuyên gia trích xuất & lọc dữ liệu.

NHIỆM VỤ:
Từ văn bản sau, hãy PHÂN TÍCH và TRÍCH XUẤT những cặp key-value quan trọng (tối đa 12 cặp)
và trả về dưới dạng JSON, ví dụ: {{"tên": "...", "số": "..."}}.

VĂN BẢN:
{text}

HƯỚNG DẪN:
1. Chỉ trả về một JSON duy nhất, không có văn bản khác.
2. Key nên ngắn gọn, tiếng Việt không dấu hoặc có dấu, tuỳ hợp lý.
3. Nếu không tìm thấy, để value là chuỗi rỗng "".
4. Trả về tối đa 12 trường.
"""
    else:
        # Tạo example JSON
        example_dict = {ph: f"[thông tin về {ph}]" for ph in placeholders}
        prompt = f"""Bạn là chuyên gia trích xuất & lọc dữ liệu.

NHIỆM VỤ:
Từ văn bản sau, hãy PHÂN TÍCH và TRÍCH XUẤT các thông tin phù hợp với các trường sau:
{', '.join(placeholders)}

VĂN BẢN:
{text}

HƯỚNG DẪN:
1. Đọc kỹ văn bản để hiểu nội dung chính
2. Lọc ra thông tin LIÊN QUAN đến các trường được yêu cầu
3. Nếu không tìm thấy thông tin nào, để value là chuỗi rỗng ""
4. Trả về CHÍNH XÁC định dạng JSON (không có text khác)
5. Giá trị phải ngắn gọn, rõ ràng

FORMAT OUTPUT (CHỈNH XÁC):
{json.dumps(example_dict, ensure_ascii=False)}
"""
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
        )
        
        response_text = response.choices[0].message.content.strip()
        
        # Try to parse JSON
        try:
            data_dict = json.loads(response_text)
            # Ensure all placeholders are in dict
            for ph in placeholders:
                if ph not in data_dict:
                    data_dict[ph] = ""
        except json.JSONDecodeError:
            # Fallback: return default empty dict
            data_dict = {ph: "" for ph in placeholders}
            st.warning(f"⚠️ Không thể parse JSON từ AI. Response: {response_text[:100]}")
        
        return data_dict
    except Exception as e:
        raise Exception(f"Lỗi gọi OpenAI: {str(e)}")




# GIAO DIỆN CHÍNH CỦA ỨNG DỤNG
# ========================================================================================

st.title("📄 Trợ lý OCR & Điền Mẫu Thông minh")

# Tạo 2 tabs: OCR và Điền Mẫu
tab1, tab2 = st.tabs(["📊 OCR (Trích xuất text)", "📝 Điền Mẫu (Template Filling)"])

with tab2:
    st.header("Điền Mẫu Tự Động")
    st.write("Upload file mẫu + file dữ liệu → App sẽ tự động điền dữ liệu vào mẫu.")
    
    col_template, col_files = st.columns(2)
    
    with col_template:
        st.subheader("1️⃣ Chọn file mẫu")
        template_file = st.file_uploader(
            "Upload file Word mẫu (.docx) - chứa placeholders dạng {tên_trường}",
            type=['docx'],
            key="template_uploader"
        )
    
    with col_files:
        st.subheader("2️⃣ Chọn file dữ liệu")
        st.info("""
        Upload các file có dữ liệu:
        - Ảnh, PDF (sẽ dùng OCR)
        - Word, Excel, Text (đọc trực tiếp)
        - Tối đa 20 file
        """)
        data_files = st.file_uploader(
            "Chọn file dữ liệu",
            type=['png', 'jpg', 'jpeg', 'pdf', 'docx', 'xlsx', 'txt'],
            accept_multiple_files=True,
            key="data_files_uploader"
        )
    
    # Khởi tạo session state
    if "extracted_data" not in st.session_state:
        st.session_state.extracted_data = {}
    if "last_uploaded_files" not in st.session_state:
        st.session_state.last_uploaded_files = None
    
    # Kiểm tra nếu file upload thay đổi → reset dữ liệu cũ
    current_file_names = tuple(f.name for f in data_files) if data_files else None
    if current_file_names != st.session_state.last_uploaded_files:
        st.session_state.extracted_data = {}
        st.session_state.last_uploaded_files = current_file_names
    
    if template_file and data_files:
        st.markdown("---")
        
        # Detect placeholders trong template
        template_bytes = template_file.getvalue()
        placeholders = detect_placeholders_in_template(template_bytes)
        
        st.subheader("3️⃣ Các trường cần điền:")
        if placeholders:
            st.success(f"Tìm thấy {len(placeholders)} trường: {', '.join(placeholders)}")
        else:
            st.warning("Không tìm thấy placeholder nào! Kiểm tra template có định dạng {tên_trường}?")
        
        st.markdown("---")
        st.subheader("4️⃣ Trích xuất thông tin từ file")
        
        col_extract, col_clear = st.columns([2, 1])
        
        with col_extract:
            if st.button("🔍 Trích xuất thông tin", type="secondary", use_container_width=True):
                st.session_state.extracted_data = {}
                progress_bar = st.progress(0, text="Đang trích xuất...")
                
                for idx, data_file in enumerate(data_files):
                    progress_bar.progress(
                        (idx + 1) / len(data_files),
                        text=f"Trích xuất {idx + 1}/{len(data_files)}: {data_file.name}"
                    )
                    
                    try:
                        # Bước 1: Trích xuất text
                        file_ext = data_file.name.split('.')[-1].lower()
                        file_bytes = data_file.getvalue()
                        extracted_text = extract_text_from_file(file_bytes, file_ext)
                        
                        if extracted_text.startswith("Lỗi"):
                            st.warning(f"⚠️ {data_file.name}: {extracted_text}")
                            continue
                        
                        # Bước 2: AI phân tích và lọc dữ liệu
                        data_dict = extract_structured_data_with_ai(extracted_text, placeholders)
                        
                        st.session_state.extracted_data[data_file.name] = {
                            "text": extracted_text,
                            "data": data_dict,
                            "file": data_file
                        }
                    
                    except Exception as e:
                        st.warning(f"⚠️ Lỗi xử lý {data_file.name}: {str(e)}")
                
                progress_bar.empty()
                st.rerun()
        
        with col_clear:
            if st.button("🗑️ Xóa dữ liệu", type="secondary", use_container_width=True):
                st.session_state.extracted_data = {}
                st.rerun()
        
        # Hiển thị dữ liệu đã trích xuất
        if "extracted_data" in st.session_state and st.session_state.extracted_data:
            st.markdown("---")
            st.subheader("5️⃣ Dữ liệu được trích xuất & lọc:")
            st.info(f"📦 Đã lưu dữ liệu từ {len(st.session_state.extracted_data)} file. Bạn có thể sửa dữ liệu bên dưới.")
            
            for file_name, file_data in st.session_state.extracted_data.items():
                with st.expander(f"📄 {file_name}", expanded=True):
                    st.write("**Thông tin được lọc:**")
                    
                    # Hiển thị dữ liệu được lọc
                    for key, value in file_data["data"].items():
                        col1, col2 = st.columns([1, 3])
                        with col1:
                            st.text(f"**{key}:**")
                        with col2:
                            # Cho phép edit thông tin
                            edited_value = st.text_input(
                                label="edit",
                                value=str(value),
                                key=f"edit_{file_name}_{key}",
                                label_visibility="collapsed"
                            )
                            file_data["data"][key] = edited_value
                    
                    # Hiển thị text gốc
                    with st.expander("📋 Text gốc trích xuất"):
                        st.text_area(
                            "Nội dung ban đầu:",
                            value=file_data["text"],
                            height=150,
                            disabled=True,
                            key=f"text_{file_name}"
                        )
            
            st.markdown("---")
            st.subheader("6️⃣ Tạo file từ mẫu")
            
            if st.button("✨ Điền mẫu và tạo file", type="primary", use_container_width=True):
                try:
                    filled_files = []
                    
                    for file_name, file_data in st.session_state.extracted_data.items():
                        try:
                            # Debug: Hiển thị dữ liệu sẽ được điền
                            st.info(f"📝 Dữ liệu sẽ điền cho {file_name}:")
                            for key, value in file_data["data"].items():
                                st.text(f"  {{{key}}} → {value}")
                            
                            # Điền vào template (dùng placeholder nếu có, ngược lại detect labels + map và dùng label-fallback)
                            if placeholders:
                                filled_bytes = fill_template_word(template_bytes, file_data["data"])
                            else:
                                # detect labels in template
                                template_labels = detect_labels_in_template(template_bytes)
                                st.info(f"🔎 Detected labels in template: {template_labels}")
                                # map AI keys to template labels
                                mapped = map_data_to_labels(file_data["data"], template_labels)
                                st.info(f"🔁 Mapping label -> value: {mapped}")
                                # use mapped values to fill labels
                                filled_bytes = fill_template_with_labels(template_bytes, mapped)
                                # update saved data for display/download
                                file_data["data_mapped"] = mapped
                            
                            output_name = file_name.rsplit('.', 1)[0]
                            filled_files.append({
                                "name": f"{output_name}_filled.docx",
                                "bytes": filled_bytes,
                                "data": file_data["data"],
                                "source": file_name
                            })
                        except Exception as e:
                            st.warning(f"⚠️ Lỗi điền mẫu cho {file_name}: {str(e)}")
                    
                    if filled_files:
                        st.success(f"✅ Đã tạo {len(filled_files)} file thành công!")
                        
                        st.markdown("---")
                        st.subheader("📥 Tải kết quả:")
                        
                        # Download từng file
                        for item in filled_files:
                            col1, col2 = st.columns([3, 1])
                            with col1:
                                st.text(f"📄 {item['name']} (từ: {item['source']})")
                            with col2:
                                st.download_button(
                                    "⬇️ Tải",
                                    data=item['bytes'],
                                    file_name=item['name'],
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"dl_filled_{item['name']}"
                                )
                        
                        # Option tải tất cả (zip)
                        st.markdown("---")
                        if st.button("📦 Tải tất cả file (zip)", use_container_width=True):
                            zip_buffer = io.BytesIO()
                            with zipfile.ZipFile(zip_buffer, 'w') as zf:
                                for item in filled_files:
                                    zf.writestr(item['name'], item['bytes'])
                            
                            zip_buffer.seek(0)
                            st.download_button(
                                "⬇️ Tải ZIP",
                                data=zip_buffer.getvalue(),
                                file_name="ket_qua_dien_mau.zip",
                                mime="application/zip",
                                key="dl_zip_all"
                            )
                    else:
                        st.error("❌ Không thể tạo file nào.")
                
                except Exception as e:
                    st.error(f"❌ Lỗi: {str(e)}")






with tab1:
    st.header("Trích xuất văn bản từ ảnh hoặc PDF")
    st.write("Mặc định xử lý song ngữ Tiếng Việt và Tiếng Anh.")

    # Cột cho phần tải lên và hướng dẫn
    col1, col2 = st.columns([2, 1])

    with col1:
        uploaded_files_all = st.file_uploader(
            f"Tải lên nhiều file (tối đa {MAX_FILES} file cùng lúc)",
            type=['pdf', 'png', 'jpg', 'jpeg'],
            accept_multiple_files=True
        )

    with col2:
        with st.expander("💡 Mẹo sử dụng", expanded=True):
            st.info(f"""
            - Tối đa **{MAX_FILES} file** mỗi lần; hỗ trợ song ngữ Việt + Anh.
            - Có thể chọn **Xử lý AI cho tất cả** sau khi OCR xong.
            - Mỗi file: TXT, Word (A4), Excel (A4); font Times New Roman.
            """)

    # Giới hạn 20 file, giữ thứ tự
    uploaded_files = list(uploaded_files_all)[:MAX_FILES] if uploaded_files_all else []
    if uploaded_files_all and len(uploaded_files_all) > MAX_FILES:
        st.warning(f"Chỉ xử lý {MAX_FILES} file đầu tiên. Tổng số file chọn: {len(uploaded_files_all)}.")

    # Khởi tạo session state cho kết quả OCR và AI
    if "ocr_results" not in st.session_state:
        st.session_state.ocr_results = []
    if "ocr_file_keys" not in st.session_state:
        st.session_state.ocr_file_keys = ()
    if "ai_results" not in st.session_state:
        st.session_state.ai_results = {}

    # Xử lý OCR hàng loạt khi có file mới hoặc đổi danh sách
    file_keys = tuple((f.name, f.size) for f in uploaded_files) if uploaded_files else ()
    if uploaded_files and file_keys != st.session_state.ocr_file_keys:
        st.session_state.ocr_file_keys = file_keys
        st.session_state.ocr_results = []
        progress_bar = st.progress(0, text="Đang OCR...")
        for idx, uf in enumerate(uploaded_files):
            progress_bar.progress((idx + 1) / len(uploaded_files), text=f"Đang xử lý file {idx + 1}/{len(uploaded_files)}: {uf.name}")
            file_bytes = uf.getvalue()
            ext = uf.name.split('.')[-1].lower()
            text, err = process_file(file_bytes, ext, show_progress=False)
            st.session_state.ocr_results.append({"name": uf.name, "text": text, "error": err})
        progress_bar.empty()
        st.session_state.ai_results = {}  # Reset AI khi đổi bộ file
        st.rerun()

    # Hiển thị kết quả từng file
    if uploaded_files and st.session_state.ocr_results:
        st.markdown("---")
        st.header("Kết quả trích xuất")

        # Nút xử lý AI cho tất cả file
        run_ai_all = st.button("✨ Xử lý AI cho tất cả các file", type="primary", use_container_width=True)
        if run_ai_all:
            try:
                bar = st.progress(0, text="Đang gọi OpenAI...")
                n = len(st.session_state.ocr_results)
                for i, res in enumerate(st.session_state.ocr_results):
                    if res["error"]:
                        continue
                    bar.progress((i + 1) / n, text=f"AI đang xử lý {i + 1}/{n}: {res['name']}")
                    fixed = call_openai_proofread(res["text"])
                    st.session_state.ai_results[res["name"]] = fixed
                bar.empty()
                st.success("Đã xử lý AI xong tất cả file.")
                st.rerun()
            except Exception as e:
                st.error(f"Lỗi khi gọi OpenAI: {e}")

        for i, uploaded_file in enumerate(uploaded_files):
            res = st.session_state.ocr_results[i] if i < len(st.session_state.ocr_results) else None
            if not res:
                continue
            name, text, error = res["name"], res["text"], res["error"]

            with st.expander(f"📄 {name}", expanded=(i < 3)):
                if error:
                    st.error(error)
                else:
                    st.text_area("Văn bản OCR:", text, height=220, key=f"text_{name}_{i}")

                    col_txt, col_docx, col_xlsx = st.columns(3)
                    with col_txt:
                        st.download_button(
                            label="📥 Tải TXT",
                            data=text.encode('utf-8'),
                            file_name=f"ket_qua_{name}.txt",
                            mime="text/plain",
                            key=f"dl_txt_{name}_{i}"
                        )
                    with col_docx:
                        st.download_button(
                            label="📄 Tải Word (A4)",
                            data=build_docx(text),
                            file_name=f"ket_qua_{name}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_docx_{name}_{i}"
                        )
                    with col_xlsx:
                        st.download_button(
                            label="📊 Tải Excel (A4)",
                            data=build_excel(text),
                            file_name=f"ket_qua_{name}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            key=f"dl_xlsx_{name}_{i}"
                        )

                    # Xử lý AI từng file
                    use_ai_one = st.button("✨ Dùng AI cho file này", key=f"ai_one_{name}_{i}")
                    if use_ai_one:
                        try:
                            with st.spinner("OpenAI đang xử lý..."):
                                fixed_text = call_openai_proofread(text)
                            st.session_state.ai_results[name] = fixed_text
                            st.rerun()
                        except Exception as e:
                            st.error(f"Lỗi OpenAI: {e}")

                    # Hiển thị kết quả AI nếu đã có
                    fixed_text = st.session_state.ai_results.get(name)
                    if fixed_text:
                        st.markdown("---")
                        st.subheader("Văn bản đã được AI hiệu đính")
                        st.text_area("", fixed_text, height=220, key=f"text_ai_{name}_{i}")

                        ai_txt, ai_docx, ai_xlsx = st.columns(3)
                        with ai_txt:
                            st.download_button(
                                label="📥 Tải TXT (AI)",
                                data=fixed_text.encode('utf-8'),
                                file_name=f"ket_qua_AI_{name}.txt",
                                mime="text/plain",
                                key=f"dl_ai_txt_{name}_{i}"
                            )
                        with ai_docx:
                            st.download_button(
                                label="📄 Tải Word (A4, AI)",
                                data=build_docx(fixed_text),
                                file_name=f"ket_qua_AI_{name}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"dl_ai_docx_{name}_{i}"
                            )
                        with ai_xlsx:
                            st.download_button(
                                label="📊 Tải Excel (A4, AI)",
                                data=build_excel(fixed_text),
                                file_name=f"ket_qua_AI_{name}.xlsx",
                                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                key=f"dl_ai_xlsx_{name}_{i}"
                            )

